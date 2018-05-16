from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_BREAK
from datetime import datetime
from officer import Officer
from officer import Title
from attendance import Attendance_API
from db import DB

class Meeting_generator:


	def __init__(self,path_to_folder, meeting_date=datetime.now()):
		document = Document()
		db = DB()
		'''you can only modify font/font size/bold through a run'''

		document.styles['Normal'].font.name = 'Times New Roman'
		document.styles['Normal'].font.size = Pt(12)
		document.styles['Normal'].paragraph_format.line_spacing = 1
		document.styles['Normal'].paragraph_format.space_after = Pt(0)
		'''heading'''
		header_paragraph = document.add_paragraph()
		header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		header_paragraph.paragraph_format.line_spacing = 1
		header_paragraph.paragraph_format.space_after = Pt(1) #WTF WHY DID THIS WORK
		header_run = header_paragraph.add_run()


		header_run.bold = True
		header_run.font.name = 'Georgia'
		header_run.font.size = Pt(16)

		header_run.add_picture('secret/Alpha_sig.png', width=Inches(1.52))
		header_run.add_break()

		'''possibly change this to work with API call'''
		header_run.add_text('ZETA GAMMA CHAPTER')
		header_run.add_break()
		'''possibly change this to work with API call'''
		header_run.add_text('UNIVERSITY OF CALIFORNIA, DAVIS')

		'''two breaks in a row doesn't work aparently?'''
		header_run.add_break(WD_BREAK.LINE)
		header_run.add_break(WD_BREAK.LINE)
		header_run.add_break(WD_BREAK.LINE)


		'''MM DD, YYYY'''
		#not sure about %d and %-d
		date_string = 'Chapter Meeting Minutes - ' + meeting_date.strftime('%B %-d, %Y')

		date_run = header_paragraph.add_run()
		date_run.bold = True
		date_run.font.size = Pt(16)

		date_run.add_text(date_string)
		date_run.add_break(WD_BREAK.LINE)
		date_run.add_break(WD_BREAK.LINE)

		'''
			body paragraph contains:

		'''
		body_paragraph = document.add_paragraph()
		body_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
		#body_paragraph.paragraph_format.line_spacing = 0	doesn't work
		body_paragraph.line_spacing_rule =  WD_LINE_SPACING.EXACTLY 
		body_paragraph.paragraph_format.space_after = Pt(0)

		'''
			meetings usually start at 5PM, which is 1700. 
			im assuming is doc is generated after 5, meeting isn't starting at 5:00 pm
		'''
		time = meeting_date.strftime('%I:%M%p')
		attendance_run = body_paragraph.add_run()
		attendance_run.font.size = Pt(12)
		attendance_run.bold = True

		attendance_run.add_text('MEETING TO ORDER: ')
		attendance_run.add_text(time)
		attendance_run.add_break(WD_BREAK.LINE)
		attendance_run.add_break(WD_BREAK.LINE)
		attendance_run.add_text('ROLL CALL: ')
		attendance_run.add_break(WD_BREAK.LINE)


		Attendance = Attendance_API()
		today_attendance = Attendance.get_attendance('05/22')
		table = document.add_table(10,3)
		table.style = 'Table Grid'
		table_cells = []
		for colIndex, col in enumerate(table.columns):
			for rowIndex, cell in enumerate(col.cells):
				text = str(colIndex * 10 + rowIndex+1)
				cell.text = text + '.'
				table_cells.append(cell)

		for (num, member) in zip(table_cells, today_attendance):
			num.text = num.text + ' ' + member[0] + ' ' + member[1]

			

		prudential_paragraph = document.add_paragraph()

		quorum_run = prudential_paragraph.add_run()
		#quorum_run.font.size = Pt(12)
		quorum_run.bold = True

		quorum_run.add_break(WD_BREAK.LINE)
		quorum_run.add_text('Quorum: Met/Not Met')

		quorum_run.add_break(WD_BREAK.LINE)
		'''
		prudential_underline = prudential_paragraph.add_run('PRUDENTIAL REPORTS')
		#prudential_underline.font.size = Pt(12)
		prudential_underline.bold = True
		prudential_underline.underline = True
		#prudential_underline.add_break(WD_BREAK.LINE)
		'''
		officer_title = Title(document, 'PRUDENTIAL REPORTS', font = 'Arial')

		officer_response = db.get_officers()
		
		officer_list = []
		for officer in officer_response:
			officer_list.append(Officer(document, officer))
			

		'''now begin actual reports'''

		director_title = Title(document, 'DIRECTOR REPORTS', font = 'Arial')

		director_response = db.get_directors()
		director_list = []
		for director in director_response:
			director_list.append(Officer(document, director))

		VP_title = Title(document, 'VICE PRESIDENT\'S REPORT', font = 'Arial', bold = True, underline = True)
		vp = Officer(document,('Vice President', 'Max'))

		president_title = Title(document, 'PRESIDENT\'S REPORT', font = 'Arial', bold = True, underline = True)
		vp = Officer(document,('President', 'Alex'))

		old_biz = 	Title(document, string = 'OLD BUSINESS', font = 'Arial', 	bold = True, underline = True)
		new_biz = 	Title(document, string = 'NEW BUSINESS', font = 'Arial', 	bold = True, underline = True)
		poll = 		Title(document, string = 'Move to Poll', font = 'Arial', 	bold = True)
		calendar = 	Title(document, string = 'Move to Amend the Calendar', font = 'Arial', bold = True, underline = False)
		call_outs = Title(document, string = 'CALL OUTS', 	font = 'Arial',		bold = True)	
		comments =  Title(document, string = 'COMMENTS FOR THE GOOD OF SOCIETY', font = 'Arial', bold = True)	

		adjourned = Title(document, string = 'MEETING ADJOURNED: Time', font = 'Arial', bold = True)	

		print 'just saved to: ', path_to_folder, meeting_date.strftime('%Y.%m.%d') + '.docx'

		document.save(path_to_folder + meeting_date.strftime('%Y.%m.%d') + '.docx')
