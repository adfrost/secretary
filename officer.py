'''
	- a level of indirection to easily iterate over all officer positions and group the two paragraphs that each require
	- otherwise, there would be no differentiation between a 'paragraph' for an officer's position and name, and the report
	
'''
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_BREAK
from datetime import datetime

class Position:

	def __init__(self, document, position):

		'''
			could make a scraper crawl through the portal, 
			or host everything locally and have it occasionally compare with the portal
		self.incumbent = get_officer_name(position)

		'''

		self.text = position[0] + ' (' + position[1] + '):'

		self.officer = document.add_paragraph(style = 'List Paragraph')
		self.officer.paragraph_format.space_after = Pt(0)
		self.officer.add_run(self.text).bold = True
		self.officer.paragraph_format.line_spacing = 1

		self.report = document.add_paragraph('Report: ', style = 'List Bullet 3')
		self.report.bold = False
		self.report.paragraph_format.space_after = Pt(0)
		self.report.paragraph_format.space_before = Pt(0)
		self.report.add_run().add_break(WD_BREAK.LINE)


'''
	Title class easily adds a bolded, underlined line with custom text
	and optional font. Title class adds a one line break before and one
	after the text

 oh look its Dennis
'''

class Title:

	def __init__(self, document, string, font = None, bold = False, underline = False):
		self.string = string

		self.p = document.add_paragraph()
		self.run = self.p.add_run()
		self.run.add_break(WD_BREAK.LINE)
		self.run.add_text(string.upper())
		self.run.bold = bold
		self.run.underline = underline
		self.run.add_break(WD_BREAK.LINE)

		self.run.font.name = font

		self.p.paragraph_format.space_before = Pt(0)
		self.p.paragraph_format.space_after = Pt(0)


'''not sure where this should go since other files now use it besides attendance.py'''
'''
	takes in datetime object
'''
def get_quarter(date=None):
	if not date:
		date = datetime.now()
	quarter = ''
	if date.month >= 9:
		quarter = 'Fall '
	elif date.month <= 3:
		quarter = 'Winter '
	elif 4 <= date.month <= 7:
		quarter = 'Spring '
	else:
		print('error, current date is in summer')
		quarter = 'ERROR '

	return quarter + str(date.year)


#summer counts as the previous year right?
def get_academic_year(date=None):
	if not date:
		date = datetime.now()
	quarter = ''
	if date.month < 9:
		quarter += str(date.year - 1) + '-' + str(date.year)
	else:
		quarter += str(date.year) + '-' + str(date.year + 1)
		
	return quarter

#should return either '[Chapter|Prudential]/'[academic year]''
#if date is Nonetype, function is called in the 'today' condition
# ex: returns "Chapter/2017-2018/2018.04.20"
def get_file_path(date=None):
	if not date:
		date = datetime.now()

	path_to_file = ''
	
	if(date.weekday() == 6):	#meetings are traditionally held on sundays
		path_to_file += 'Chapter/'
	else:
		path_to_file += 'Prudential/'

	path_to_file += get_academic_year(date) + '/'
	if not date:			
		path_to_file += date.strftime('%Y.%m.%d')
	

	return path_to_file

